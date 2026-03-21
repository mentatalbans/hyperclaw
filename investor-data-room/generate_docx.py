#!/usr/bin/env python3
"""Generate .docx files for all HyperClaw investor data room documents."""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import html

BASE_DIR = Path("/Users/mentat/.openclaw/workspace/hyperclaw/investor-data-room")

# Document content map (title, sections)
DOCS = [
    {
        "path": "00-index/HC-Data-Room-Index.html",
        "title": "HyperClaw — Investor Data Room Index",
        "sections": [
            ("HyperClaw AI Orchestration Platform", "Investor Data Room — Q1 2026\nConfidential & Proprietary\nPre-Seed Round: $2M at $10M Pre-Money Valuation"),
            ("About This Data Room", "This investor data room contains all materials necessary to evaluate a pre-seed investment in HyperClaw. Documents are organized across 8 sections covering company overview, pitch materials, financial projections, market analysis, product architecture, team background, legal/risk factors, and traction data."),
            ("Section 00 — Index", "HC-Data-Room-Index.html / .pdf / .docx\nThis document."),
            ("Section 01 — Company Overview", "HC-Company-One-Pager — At-a-glance company summary\nHC-Executive-Summary — Detailed executive overview"),
            ("Section 02 — Pitch Deck", "HC-Investor-Pitch-Deck — Full investor presentation"),
            ("Section 03 — Financials", "HC-Cap-Table — Capitalization table\nHC-Financial-Model — 5-year financial projections\nHC-Use-of-Funds — Pre-seed capital allocation"),
            ("Section 04 — Market", "HC-Market-Sizing — TAM/SAM/SOM analysis ($47B opportunity)\nHC-Competitive-Landscape — Feature matrix vs LangChain, AutoGen, CrewAI"),
            ("Section 05 — Product", "HC-Product-Overview — 22 agents, 4 pillars, enterprise tiers\nHC-Technology-Architecture — Deep technical architecture"),
            ("Section 06 — Team", "HC-Leadership-Team — Vaughn Davis (CEO) & Andy Zhang (CTO)"),
            ("Section 07 — Legal", "HC-Risk-Factors — 7 risk factors with mitigants"),
            ("Section 08 — Traction", "HC-Traction-Summary — v0.1 alpha, 377 tests, $2M raise details"),
            ("Contact", "GitHub: github.com/mentatalbans/hyperclaw\nLicense: MIT (open-source core)\nClassification: Confidential — Do Not Distribute"),
        ]
    },
    {
        "path": "01-overview/HC-Company-One-Pager.html",
        "title": "HyperClaw — Company One-Pager",
        "sections": [
            ("HyperClaw — Company Overview", "The Production-Grade Multi-Agent AI Orchestration Platform\nQ1 2026 | Confidential"),
            ("The Problem", "Enterprise AI agents fail in production. Open-source frameworks like LangChain, AutoGen, and CrewAI were built for research and experimentation, not enterprise deployment. They lack policy enforcement, intelligent routing, production-grade memory, and typed state management.\n\nThe result: 73% of enterprise AI projects fail to reach production. The missing layer is orchestration infrastructure built for scale."),
            ("The Solution — HyperClaw", "HyperClaw is the first AI orchestration platform designed for enterprise production from day one. Four architectural pillars:\n\n• HyperState — Pydantic v2 typed state management. Every interaction serializable and auditable.\n• HyperRouter — UCB1 bandit algorithm. Intelligent routing to optimal agents based on performance.\n• HyperMemory — pgvector causal graph. Agents reason across temporal dependencies.\n• HyperShield — Policy engine. Pre-execution action validation for compliance teams.\n\nAbove them all: PROMETHEUS — the recursive meta-agent that orchestrates the entire system."),
            ("22 Specialist Agents", "Personal: LifeCoach, HealthAdvisor, FinanceAdvisor, TravelPlanner, LegalAdvisor\nBusiness: SalesIntelligence, MarketingStrategist, OperationsOptimizer, CustomerSuccess, HRAdvisor, ExecutiveAssistant\nScientific: ResearchAnalyst, DataScientist, BioinformaticsAgent, ClimateModeler\nCreative: ContentStrategist, NarrativeWriter, BrandVoice, MediaProducer\nRecursive: PROMETHEUS, AgentBuilder, OrchestrationManager"),
            ("Market Opportunity", "Total Addressable Market: $47B (AI agent infrastructure + enterprise software)\nServing the enterprise AI orchestration segment, estimated $8.2B SAM growing at 34% CAGR.\nHyperClaw's initial SOM: $420M targeting mid-market to enterprise buyers."),
            ("Business Model", "Starter: $499/mo — Teams exploring AI automation\nPro: $1,999/mo — Growing enterprises\nEnterprise: $9,999/mo — Large organizations with compliance needs\nSovereign: Custom — Air-gapped, government deployments"),
            ("Traction", "v0.1.0-alpha: Live on GitHub (github.com/mentatalbans/hyperclaw)\n377 tests passing, zero failures\nSupabase backend live\nMIT licensed open-source core\nPre-seed raise: $2M at $10M pre-money"),
            ("Team", "Vaughn Davis — Founder & CEO: 15+ years luxury hospitality, former MD Dream Hollywood, CEO Hyper Nimbus & Hyper Talent\nAndy Zhang — Co-Founder & CTO: Co-Founder Fairsplit, Co-Founder Cogniscale, AI infrastructure specialist"),
            ("The Ask", "Raising: $2,000,000\nPre-money valuation: $10,000,000\nUse: 40% engineering hires, 25% GTM & sales, 20% infrastructure, 15% operations/legal\nARR targets: $500K (2026) → $3M (2027) → $12M (2028) → $50M (2030)"),
        ]
    },
    {
        "path": "01-overview/HC-Executive-Summary.html",
        "title": "HyperClaw — Executive Summary",
        "sections": [
            ("HyperClaw Executive Summary", "Pre-Seed Investment Opportunity | Q1 2026 | Confidential"),
            ("Company Overview", "HyperClaw is a production-grade multi-agent AI orchestration platform. We build the infrastructure layer that makes enterprise AI agents actually work — with intelligent routing, causal memory, policy enforcement, and typed state management built in from day one.\n\nFounded in 2025, HyperClaw has shipped v0.1.0-alpha with 22 specialist agents across 5 domains, 377 passing tests, and a live Supabase backend. The MIT-licensed open-source core is live on GitHub at github.com/mentatalbans/hyperclaw."),
            ("The Problem We Solve", "The AI agent market is massive and growing rapidly, but 73% of enterprise AI projects fail to reach production. The failure mode is always the same: frameworks built for research (LangChain, AutoGen, CrewAI) don't have the architectural foundations enterprises need.\n\nWhat's missing:\n• Intelligent routing (not manual chains)\n• Causal memory (not just buffers)\n• Policy enforcement (not bolted-on guardrails)\n• Typed state management (not silent failures)\n\nHyperClaw is the platform built for enterprises, not developers experimenting."),
            ("The HyperClaw Architecture", "Four production-grade pillars:\n\n1. HyperState — Pydantic v2 typed, validated, immutable state management. Every agent interaction is fully serializable, auditable, and reproducible.\n\n2. HyperRouter — UCB1 multi-armed bandit algorithm selects optimal agents based on historical performance. Gets smarter with every task.\n\n3. HyperMemory — pgvector causal graph memory. Agents reason across temporal dependencies — not just what happened, but why.\n\n4. HyperShield — Pre-execution policy enforcement engine. Every agent action validated against configurable rules before it executes.\n\nPROMETHEUS sits above all four as a recursive meta-agent — the executive intelligence that spawns, coordinates, and synthesizes results from multi-agent workflows."),
            ("Market Opportunity", "Total Addressable Market: $47B across AI agent infrastructure, enterprise software, and the open-source developer economy.\n\nServing the enterprise AI orchestration segment ($8.2B SAM at 34% CAGR). Key market drivers: AI Act compliance requirements, proliferation of LLM APIs, shift from AI demos to AI deployment.\n\nHyperClaw's initial SOM: $420M in mid-market and enterprise accounts."),
            ("Competitive Advantage", "HyperClaw wins where competitors are architecturally limited:\n\n• UCB1 intelligent routing — no competitor ships adaptive bandit routing\n• HyperMemory causal graph — competitors have buffers, not causal reasoning\n• HyperShield policy engine — no competitor ships a dedicated policy engine\n• 22 specialist agents OOTB — competitors ship templates; customers build everything\n• Enterprise SLA available — no OSS competitor offers this\n• MIT open core — community adoption as the top-of-funnel moat"),
            ("Business Model & Projections", "Four tiers: Starter ($499/mo), Pro ($1,999/mo), Enterprise ($9,999/mo), Sovereign (custom)\n\nARR Targets:\n• 2026: $500,000\n• 2027: $3,000,000\n• 2028: $12,000,000\n• 2030: $50,000,000\n\nUnit economics improve significantly at scale as infrastructure costs are largely fixed and LLM costs are passed through to customers."),
            ("Leadership Team", "Vaughn Davis — Founder & CEO: 15+ years luxury lifestyle hospitality, former Managing Director of Dream Hollywood, CEO of Hyper Nimbus and Hyper Talent. Board member: Hollywood Chamber of Commerce, The Hollywood Partnership. Pioneers AI deployment in enterprise environments.\n\nAndy Zhang — Co-Founder & CTO: Co-Founder of Fairsplit (fintech) and Cogniscale (AI infrastructure). Designed and built all four HyperClaw architectural pillars and the PROMETHEUS intelligence layer. 377-test codebase reflects engineering rigor."),
            ("Investment Terms", "Pre-seed raise: $2,000,000\nPre-money valuation: $10,000,000\nPost-money: $12,000,000\nProjeted runway: 18 months\nTarget close: Q2 2026\n\nCapital use: Engineering hires (40%), GTM & sales (25%), Infrastructure (20%), Operations (10%), Legal (5%)"),
            ("Why Now", "The AI agent market is at an inflection point. LLM APIs have commoditized inference; the value is shifting to orchestration. Enterprise buyers have moved past AI experimentation and are demanding production-grade infrastructure. HyperClaw is positioned at the exact moment when the enterprise market needs what we've built."),
        ]
    },
    {
        "path": "02-pitch-deck/HC-Investor-Pitch-Deck.html",
        "title": "HyperClaw — Investor Pitch Deck",
        "sections": [
            ("HyperClaw Investor Pitch Deck", "Pre-Seed Round — $2M at $10M Pre-Money\nQ1 2026 | Confidential"),
            ("Slide 1: Title", "HyperClaw\nThe Production-Grade Multi-Agent AI Orchestration Platform\n\nVaughn Davis, CEO | Andy Zhang, CTO\nPre-Seed | $2M | Q1 2026"),
            ("Slide 2: The Problem", "73% of enterprise AI projects fail to reach production.\n\nThe reason: research frameworks in production environments.\n\nLangChain, AutoGen, CrewAI were built for experimentation — not enterprise deployment. They have no policy engine, no intelligent routing, no causal memory, no typed state.\n\nEnterprises need infrastructure built for production from day one."),
            ("Slide 3: The Solution", "HyperClaw — built for enterprise AI deployment.\n\n4 architectural pillars:\n• HyperState: Pydantic v2 typed state — auditable, reproducible\n• HyperRouter: UCB1 bandit routing — gets smarter over time\n• HyperMemory: pgvector causal graph — reason across time\n• HyperShield: Policy engine — compliance-ready from launch\n\nPROMETHEUS: recursive meta-agent that orchestrates it all."),
            ("Slide 4: Product", "22 specialist agents. 5 domains. Ready to deploy.\n\nPersonal | Business | Scientific | Creative | Recursive\n\nEnterprise tiers:\n$499/mo Starter → $1,999/mo Pro → $9,999/mo Enterprise → Custom Sovereign"),
            ("Slide 5: Market Opportunity", "$47B Total Addressable Market\n$8.2B SAM — enterprise AI orchestration (34% CAGR)\n$420M SOM — initial mid-market/enterprise target\n\nMarket drivers: AI Act compliance, LLM proliferation, production AI shift"),
            ("Slide 6: Competitive Landscape", "vs. LangChain: No policy engine, no causal memory, no UCB1 routing\nvs. AutoGen: Research-focused, ~6 examples, no enterprise SLA\nvs. CrewAI: Role templates, no memory graph, no policy enforcement\n\nHyperClaw wins: UCB1 routing + HyperMemory + HyperShield + 22 agents OOTB"),
            ("Slide 7: Traction", "v0.1.0-alpha: Live — github.com/mentatalbans/hyperclaw\n377 tests passing. Zero failures.\nSupabase backend live.\nMIT license. Public repo.\nPre-seed raise open."),
            ("Slide 8: Business Model", "Land with open source. Convert to paid.\n\nStarter ($499/mo): Teams → Pro ($1,999/mo): Scaling → Enterprise ($9,999/mo): Compliance → Sovereign (custom): Gov/Defense\n\nNet Revenue Retention target: 130%+ (upsell to higher tiers)"),
            ("Slide 9: Financial Projections", "ARR Trajectory:\n2026: $500K (first enterprise customers)\n2027: $3M (GTM scaling)\n2028: $12M (Series A growth)\n2030: $50M (market leadership)\n\nBurn: ~$111K/month at full team\nRunway: 18 months on $2M raise"),
            ("Slide 10: Team", "Vaughn Davis — CEO\n15+ years luxury hospitality, former MD Dream Hollywood\nCEO Hyper Nimbus, CEO Hyper Talent\nBoard: Hollywood Chamber of Commerce, The Hollywood Partnership\n\nAndy Zhang — CTO\nCo-Founder Fairsplit, Co-Founder Cogniscale\nAI infrastructure specialist, HyperClaw architect"),
            ("Slide 11: The Ask", "Raising: $2,000,000\nPre-money: $10,000,000\nTarget close: Q2 2026\n\nFunds: Engineering hires (40%) + GTM (25%) + Infrastructure (20%) + Operations/Legal (15%)\n\nNext milestone: First 5 enterprise customers, $500K ARR, Series A in 18 months"),
        ]
    },
    {
        "path": "03-financials/HC-Cap-Table.html",
        "title": "HyperClaw — Capitalization Table",
        "sections": [
            ("HyperClaw Cap Table", "Pre-Seed Round | Q1 2026 | Confidential"),
            ("Pre-Seed Capitalization Summary", "This document represents the anticipated capitalization structure following the close of HyperClaw's $2M pre-seed round at a $10M pre-money valuation.\n\nAll figures are illustrative pending final legal documentation. Actual terms may vary based on investor negotiations and legal counsel."),
            ("Founding Shares", "Vaughn Davis (CEO & Founder): 4,000,000 shares — 40.0% (post-round)\nAndy Zhang (CTO & Co-Founder): 3,500,000 shares — 35.0% (post-round)\nTotal Founding Shares: 7,500,000 — 75.0%\n\nFounder vesting: 4-year vesting with 1-year cliff, standard terms."),
            ("Employee Stock Option Pool", "Option Pool: 750,000 shares — 7.5% (post-round)\nReserved for engineering hires and future key employees.\nPool to be established at close."),
            ("Pre-Seed Investors", "Pre-Seed Round: 1,666,667 shares — 16.7% (post-round)\nRaise: $2,000,000 at $1.20/share\nPre-money valuation: $10,000,000\nPost-money valuation: $12,000,000"),
            ("Advisors", "Advisor Shares: 83,333 shares — 0.8% (post-round)\nReserved for strategic advisors. 2-year vest, monthly, no cliff."),
            ("Fully Diluted Cap Table Summary", "Total Authorized Shares: 10,000,000\nFounders: 7,500,000 (75.0%)\nOption Pool: 750,000 (7.5%)\nPre-Seed Investors: 1,666,667 (16.7%)\nAdvisors: 83,333 (0.8%)\nTotal Outstanding: 10,000,000 (100%)"),
            ("Key Terms", "Instrument: Common shares (founders/employees), SAFE or preferred (investors)\nPre-money valuation: $10,000,000\nPost-money valuation: $12,000,000\nPrice per share: $1.20 (effective)\nLiquidation preference: 1x non-participating (preferred shares)\nAnti-dilution: Broad-based weighted average"),
            ("Projected Series A Dilution", "Assuming Series A of $8M at $40M pre-money (~20% dilution):\nFounders: ~60% combined\nSeries A Investors: ~20%\nPre-Seed Investors: ~13%\nOption Pool: ~6%\nAdvisors: ~1%"),
        ]
    },
    {
        "path": "03-financials/HC-Financial-Model.html",
        "title": "HyperClaw — Financial Model",
        "sections": [
            ("HyperClaw Financial Model", "5-Year Projections | Q1 2026 | Confidential"),
            ("Financial Model Assumptions", "Revenue model: SaaS subscription tiers ($499 Starter, $1,999 Pro, $9,999 Enterprise, Custom Sovereign)\nCustomer acquisition: Bottom-up via open-source, converted to paid by sales team\nChurn assumption: 5% annual (improving to 3% by Year 3)\nNet Revenue Retention: 115% Year 1, growing to 130% by Year 3\nGross margin: 70% Year 1, improving to 82% by Year 5\nHeadcount: 2 FTEs at raise → 5 FTEs (Q3 2026) → 15 FTEs (2027) → 35 FTEs (2028)"),
            ("Year 2026 Projections", "ARR Target: $500,000\nMonthly Recurring Revenue (exit): $41,667\nCustomer Count: ~25–40 (mix of Starter/Pro)\nGross Revenue: ~$350,000 (partial year)\nOperating Expenses: ~$1,800,000 (burn includes team buildout)\nNet Loss: (~$1,450,000)\nCash Position (year end): ~$550,000"),
            ("Year 2027 Projections", "ARR Target: $3,000,000\nMonthly Recurring Revenue (exit): $250,000\nCustomer Count: ~120–180\nGross Revenue: ~$2,200,000\nGross Profit: ~$1,760,000 (80% margin)\nOperating Expenses: ~$3,500,000\nNet Loss: (~$1,740,000)\nNote: Series A targeted in this window"),
            ("Year 2028 Projections", "ARR Target: $12,000,000\nMonthly Recurring Revenue (exit): $1,000,000\nCustomer Count: ~400–600\nGross Revenue: ~$9,000,000\nGross Profit: ~$7,380,000 (82% margin)\nOperating Expenses: ~$7,000,000\nApproaching EBITDA breakeven"),
            ("Year 2030 Projections", "ARR Target: $50,000,000\nMonthly Recurring Revenue: $4,166,667\nCustomer Count: 1,500–2,500+\nGross Revenue: ~$45,000,000\nGross Profit: ~$37,800,000 (84% margin)\nEBITDA Positive: Yes"),
            ("Unit Economics", "Customer Acquisition Cost (CAC) — Year 1: $2,500 (bottoms up from OSS)\nLifetime Value (LTV) — Starter tier: $14,970 (avg 30 months, 5% churn)\nLTV/CAC ratio: ~6x\nPayback period: ~8 months\n\nEnterprise tier LTV/CAC improves significantly — target 15x+ by Year 3."),
            ("Revenue by Tier (2027)", "Starter ($499/mo): 60% of customers, 25% of ARR\nPro ($1,999/mo): 30% of customers, 40% of ARR\nEnterprise ($9,999/mo): 8% of customers, 30% of ARR\nSovereign (custom): 2% of customers, 5% of ARR"),
            ("Key Milestones to Series A", "1. $500K ARR by Q4 2026 (validates pricing and GTM)\n2. 3+ enterprise logos ($9,999/mo tier)\n3. Net Revenue Retention >115%\n4. 15+ employees (engineering + GTM)\n5. v0.2 with SSO/SAML shipped\n\nExpected Series A: Q1-Q2 2027 at $40M+ pre-money"),
        ]
    },
    {
        "path": "03-financials/HC-Use-of-Funds.html",
        "title": "HyperClaw — Use of Funds",
        "sections": [
            ("HyperClaw — Use of Pre-Seed Funds", "Allocation of $2,000,000 Pre-Seed Round | Q1 2026 | Confidential"),
            ("Summary", "HyperClaw is raising $2,000,000 at a $10M pre-money valuation. The capital will be deployed over approximately 18 months to achieve $500K ARR, close first enterprise customers, and position the company for a Series A raise in Q1-Q2 2027."),
            ("Engineering & Product — 40% ($800,000)", "3 senior engineering hires (full-stack AI/backend)\nTarget hires: Senior Backend Engineer, AI/ML Engineer, DevOps/Infrastructure\nTimeline: Hires Q2-Q3 2026\nRationale: Current 2-person team cannot support enterprise feature roadmap (SSO, custom policies, PROMETHEUS v2) and concurrent customer support at scale."),
            ("Go-To-Market & Sales — 25% ($500,000)", "Sales Development Representative hire\nEnterprise sales infrastructure (CRM, sequences, demos)\nConference presence and developer community investment\nContent marketing and open-source community building\nRationale: Vaughn's network provides warm leads for Q1-Q2; structured GTM required for $3M ARR target by EOY 2027."),
            ("Infrastructure & Cloud — 20% ($400,000)", "Supabase scaling (Postgres, pgvector, Auth)\nLLM API costs (OpenAI, Anthropic, Gemini) for hosted tier\nDevelopment and staging environments\nSecurity tooling, monitoring, observability\nRationale: Supabase free tier will not support enterprise customer volume; production infrastructure requires investment before customer onboarding."),
            ("Operations — 10% ($200,000)", "Office/co-working space (optional remote-first)\nHR and recruiting costs for 3 engineering hires\nTeam travel for enterprise sales and conferences\nEquipment and tooling for new hires"),
            ("Legal & Admin — 5% ($100,000)", "Corporate legal (Delaware C-Corp formation if needed)\nIP counsel for open-source licensing strategy\nEmployee contracts and option plan documentation\nGeneral legal retainer"),
            ("18-Month Milestone Map", "Month 1-3: Close raise, make first 2 engineering hires, ship v0.1.1\nMonth 4-6: First 5 enterprise pilots, hire SDR, ship v0.2 (SSO)\nMonth 7-9: Convert 3 pilots to paid, $200K ARR\nMonth 10-12: $350K ARR, 3rd engineering hire, PROMETHEUS v2\nMonth 13-15: $500K ARR, Series A preparation\nMonth 16-18: Series A close at $40M+ pre-money"),
        ]
    },
    {
        "path": "04-market/HC-Market-Sizing.html",
        "title": "HyperClaw — Market Sizing",
        "sections": [
            ("HyperClaw Market Sizing", "$47B Market Opportunity in AI Orchestration | Q1 2026 | Confidential"),
            ("Market Overview", "HyperClaw operates at the intersection of AI agent infrastructure, enterprise software, and the open-source developer economy. The market is early, large, and winner-take-most.\n\nTotal Addressable Market: $47B\nHyperClaw's opportunity spans the full enterprise AI orchestration stack — from developer tooling to enterprise platform licensing."),
            ("TAM — $47B Total Addressable Market", "The TAM represents the full global market for AI agent infrastructure, enterprise AI platforms, and developer tooling combined.\n\nKey segments:\n• AI Agent Infrastructure: $18B (2025 estimate, growing at 45% CAGR)\n• Enterprise AI Platforms: $22B (SaaS licensing, professional services)\n• Developer Tools & Orchestration: $7B (IDE plugins, frameworks, APIs)\n\nSources: Gartner, IDC, Forrester AI enterprise spending reports (2024-2025)"),
            ("SAM — $8.2B Serviceable Addressable Market", "HyperClaw's SAM is the enterprise AI orchestration segment specifically — companies actively deploying multi-agent AI workflows in production environments.\n\nGrowth rate: 34% CAGR (2025-2030)\nKey buyers: Fortune 1000, growth-stage tech companies, government agencies\nPrimary use cases: Operations automation, research pipelines, compliance-critical AI deployment\n\nHyperClaw is positioned to capture significant SAM share with its policy-first, enterprise-grade architecture."),
            ("SOM — $420M Initial Serviceable Obtainable Market", "HyperClaw's initial SOM focuses on mid-market to enterprise accounts where:\n• Annual AI software budget exceeds $50K\n• Compliance or policy requirements exist\n• Multi-agent workflows are actively evaluated\n• Vaughn Davis's network provides warm access\n\nInitial target: 500-1,000 enterprise accounts in Years 1-3"),
            ("Market Drivers", "1. AI Agent Proliferation: LLM APIs have commoditized inference; the value is shifting to orchestration infrastructure.\n\n2. Enterprise Compliance Pressure: EU AI Act, US executive orders, and SOC2/ISO requirements are creating demand for policy-enforced AI deployment.\n\n3. Production AI Failure: 73% enterprise AI failure rate is creating demand for purpose-built production infrastructure.\n\n4. Multi-Model Competition: GPT-4o, Claude 3.5, Gemini 2.0 compete on capability; enterprises need routing layers to select optimally.\n\n5. Open Source to Enterprise: The LangChain/CrewAI community is generating enterprise demand that OSS tools cannot serve."),
            ("Ideal Customer Profile", "Tier 1 — Enterprise Tech Company:\n• 500-5,000 employees\n• Active ML/AI team\n• Existing LLM spend >$10K/mo\n• Compliance requirements (SOC2, HIPAA, FedRAMP)\n\nTier 2 — Growth-Stage AI Company:\n• 50-500 employees\n• AI-native product building on LLM APIs\n• Need for multi-agent orchestration\n• Developer-led evaluation → enterprise procurement\n\nTier 3 — Government/Defense:\n• Sovereign tier candidates\n• Air-gapped deployment requirements\n• Long sales cycle, high ACV"),
            ("Competitive Timing", "HyperClaw enters the market at the optimal moment:\n• LangChain/AutoGen/CrewAI have proven the demand but cannot serve enterprise\n• Enterprise buyers are actively evaluating alternatives to OSS frameworks\n• The policy/compliance requirement is new (post-EU AI Act)\n• No well-funded enterprise-first orchestration platform exists yet\n\nThis is the Snowflake moment for AI orchestration."),
        ]
    },
    {
        "path": "04-market/HC-Competitive-Landscape.html",
        "title": "HyperClaw — Competitive Landscape",
        "sections": [
            ("HyperClaw Competitive Landscape", "Feature Matrix vs. LangChain, AutoGen, CrewAI | Q1 2026 | Confidential"),
            ("Market Position", "HyperClaw enters a rapidly growing AI orchestration market against open-source frameworks that were built for research, not enterprise. We are the first platform designed for production-grade multi-agent workflows from day one."),
            ("Feature Matrix", "Capability | HyperClaw | LangChain | AutoGen | CrewAI\n\nIntelligent Agent Routing: UCB1 Bandit [WIN] | Manual chains | Round-robin | Role-based\nMemory Architecture: HyperMemory Causal Graph [WIN] | Buffer/vector | Conversation buffer | Limited\nPolicy/Safety Engine: HyperShield [WIN] | None | Basic guardrails | None\nSpecialist Agents OOTB: 22 agents (5 domains) [WIN] | Templates only | ~6 examples | Role templates\nState Management: HyperState (Pydantic v2) | Partial | Basic | Partial\nEnterprise Auth/SSO: Roadmap Q2 2026 | None | None | None\nMulti-Model Routing: Native | Via LLMs | Yes | Limited\nRecursive/Meta-Agents: PROMETHEUS layer | None | Nested agents | None\nManaged Cloud Offering: Supabase-backed | LangSmith (paid) | None | None\nTest Coverage: 377 passing [WIN] | Varies | Varies | Partial\nEnterprise SLA: Available | None | None | None"),
            ("Our Four Unfair Advantages", "1. UCB1 Intelligent Routing\nHyperRouter uses Upper Confidence Bound (UCB1) multi-armed bandit algorithms to route tasks to the optimal agent based on historical performance — not manual configuration. Competitors rely on static chains or simple role assignments.\n\n2. HyperMemory Causal Graph\npgvector-backed causal memory graph that tracks not just what happened, but why — enabling agents to reason across temporal dependencies. LangChain's buffer and CrewAI's lack of persistent memory are no match.\n\n3. HyperShield Policy Engine\nEnterprise-grade policy enforcement layer that validates agent actions against configurable rules before execution. No competitor ships a dedicated policy engine.\n\n4. 22 Specialist Agents OOTB\nAcross 5 domains (Personal, Business, Scientific, Creative, Recursive), HyperClaw ships production-ready specialist agents. Competitors ship templates — customers must build everything themselves."),
            ("Pricing Comparison", "HyperClaw: $499/mo Starter (22 agents, HyperShield, UCB1 routing, Supabase cloud, Enterprise SLA)\nLangChain: $0 OSS + LangSmith $39+/mo (framework only, no agents, no policy engine)\nAutoGen: $0 OSS/Microsoft (research focus, ~6 agent examples, no cloud offering)\nCrewAI: $0 OSS + paid tiers (role templates, no memory graph, no policy engine)"),
            ("GitHub Stars Trajectory", "LangChain: ~90,000 stars (multi-year lead)\nAutoGen: ~35,000 stars\nCrewAI: ~25,000 stars\nHyperClaw: v0.1 Alpha — early stage, but enterprise architecture is the moat, not star count\n\nComparison: HyperClaw holds the same position vs. LangChain that Snowflake held vs. Hadoop in 2012. The incumbent has mindshare; the challenger has architecture."),
            ("Enterprise Readiness Comparison", "HyperClaw: Policy Engine ✓ | Audit Logging ✓ | Enterprise SLA ✓ | SSO Roadmap | Dedicated Support ✓ | Custom Deployment ✓\nLangChain: Policy Engine ✗ | Audit Logging (LangSmith) | Enterprise SLA ✗ | SSO ✗ | Support (paid plans) | Custom Deployment (DIY)\nAutoGen/CrewAI: None of the above"),
        ]
    },
    {
        "path": "05-product/HC-Product-Overview.html",
        "title": "HyperClaw — Product Overview",
        "sections": [
            ("HyperClaw Product Overview", "22 Agents. 4 Pillars. One Enterprise Platform. | Q1 2026 | Confidential"),
            ("Product Summary", "HyperClaw is a production-grade multi-agent AI orchestration platform. 22 specialist agents. 4 architectural pillars. One enterprise-ready system built for real-world deployment from day one."),
            ("22 Specialist Agents — 5 Domains", "Personal (5 agents):\n• LifeCoach — Personal development, goal tracking, habit formation\n• HealthAdvisor — Wellness guidance, symptom analysis, medical research\n• FinanceAdvisor — Budget analysis, investment guidance, tax planning\n• TravelPlanner — Itinerary creation, booking optimization, logistics\n• LegalAdvisor — Contract review, legal research, compliance guidance\n\nBusiness (6 agents):\n• SalesIntelligence — Prospect research, pipeline analysis, competitive intel\n• MarketingStrategist — Campaign planning, audience segmentation, content strategy\n• OperationsOptimizer — Process analysis, efficiency recommendations, workflow automation\n• CustomerSuccess — Churn prediction, expansion playbooks, health scoring\n• HRAdvisor — Recruiting support, performance frameworks, policy guidance\n• ExecutiveAssistant — Scheduling, research, communications management\n\nScientific (4 agents):\n• ResearchAnalyst — Literature review, hypothesis generation, citation management\n• DataScientist — Statistical analysis, model selection, visualization\n• BioinformaticsAgent — Genomic data analysis, pathway mapping\n• ClimateModeler — Environmental modeling, emissions analysis\n\nCreative (4 agents):\n• ContentStrategist — Editorial planning, SEO strategy, content calendars\n• NarrativeWriter — Long-form content, storytelling, brand narrative\n• BrandVoice — Voice consistency, tone analysis, brand guideline enforcement\n• MediaProducer — Script writing, production planning, multimedia strategy\n\nRecursive (3 agents):\n• PROMETHEUS — Recursive meta-agent, dynamic orchestration\n• AgentBuilder — Creates new specialized agents on demand\n• OrchestrationManager — Workflow coordination, resource allocation"),
            ("4 Architecture Pillars", "01. HyperState (State Management)\nTyped, validated, immutable state management built on Pydantic v2. Every agent interaction is fully serializable, auditable, and reproducible. No silent failures, no mysterious state mutations.\nTech: Pydantic v2 · Supabase · Event sourcing\n\n02. HyperRouter (Intelligent Routing)\nUCB1 multi-armed bandit algorithm routes tasks to the highest-performing agent based on historical success rates. Gets smarter over time — no manual chain configuration required.\nTech: UCB1 bandit · Performance tracking · Auto-optimization\n\n03. HyperMemory (Persistent Memory)\nCausal graph memory layer powered by pgvector. Agents remember not just facts, but causal relationships — enabling temporal reasoning across multi-session, multi-agent workflows.\nTech: pgvector · Causal graphs · Semantic search\n\n04. HyperShield (Safety & Policy)\nPre-execution policy enforcement engine. Every agent action is validated against configurable rules before it executes. Designed for compliance teams and enterprise security requirements.\nTech: Policy rules · Audit trail · Compliance-ready"),
            ("PROMETHEUS — Recursive Intelligence Layer", "PROMETHEUS is HyperClaw's recursive meta-agent, sitting above all four pillars. It can spawn, coordinate, and terminate other agents dynamically, enabling self-organizing agent networks that adapt to complex, open-ended tasks.\n\nKey capabilities:\n• Dynamic agent spawning\n• Task decomposition and delegation\n• Parallel execution trees\n• Result synthesis\n• Self-termination when objectives are met"),
            ("Enterprise Pricing Tiers", "Starter — $499/mo\nFor teams exploring AI automation. All 22 agents, HyperShield basic, 5 concurrent workflows, community support, Supabase cloud backend.\n\nPro — $1,999/mo\nFor growing enterprises. All Starter features, 25 concurrent workflows, HyperShield advanced, PROMETHEUS access, priority support + SLA, custom agent configuration.\n\nEnterprise — $9,999/mo\nFor large organizations. All Pro features, unlimited workflows, SSO/SAML (Q2 2026), custom policy engine, dedicated CSM, on-prem deployment option.\n\nSovereign — Custom pricing\nAir-gapped, sovereign deployments for government and defense. Air-gapped deployment, custom model hosting, full source access, embedded team support, classified workload capability."),
            ("Primary Use Cases", "Enterprise Automation: Specialized agent pipelines for sales intelligence, HR workflows, operations optimization, and customer success with policy-enforced guardrails.\n\nResearch & Analysis: Scientific and data science agents perform literature reviews, dataset analysis, and hypothesis generation with causal memory across sessions.\n\nRecursive Orchestration: PROMETHEUS enables self-organizing agent networks that decompose complex goals into sub-tasks, spawn specialist agents, and synthesize results autonomously.\n\nCompliance-Critical AI: HyperShield makes HyperClaw the first AI orchestration platform compliance officers can approve — with full audit trails and configurable policy rules.\n\nContent at Scale: Creative domain agents handle brand-consistent content strategy, narrative writing, and media production at enterprise volume.\n\nMulti-Model Workflows: HyperRouter selects the right LLM for each subtask — GPT-4o for reasoning, Claude for writing, Gemini for multimodal — optimizing cost and quality automatically."),
        ]
    },
    {
        "path": "05-product/HC-Technology-Architecture.html",
        "title": "HyperClaw — Technology Architecture",
        "sections": [
            ("HyperClaw Technology Architecture", "Deep Technical Reference | Q1 2026 | Confidential"),
            ("Architecture Overview", "HyperClaw is built on four production-grade subsystems — HyperState, HyperRouter, HyperMemory, and HyperShield — with PROMETHEUS as the recursive intelligence layer orchestrating above them all. 377 tests. MIT core. Zero compromises."),
            ("System Architecture — Layered Stack", "Layer 1: PROMETHEUS (Top)\nRecursive meta-agent. Spawns, coordinates, and terminates agents dynamically. Self-organizing intelligence layer for complex open-ended tasks.\nTech: Meta-agent · Recursive orchestration\n\nLayer 2: HyperShield\nPre-execution policy enforcement. Every agent action validated against configurable rules before execution. Full audit trail.\nTech: Policy engine · Audit log · Compliance\n\nLayer 3: HyperRouter\nUCB1 multi-armed bandit routing. Selects optimal agent based on historical performance. Gets smarter with every task.\nTech: UCB1 bandit · Performance tracking\n\nLayer 4: HyperMemory\nCausal graph memory powered by pgvector. Tracks facts and causal relationships — enabling temporal reasoning across multi-session workflows.\nTech: pgvector · Causal graph · Semantic search\n\nLayer 5: HyperState\nTyped, validated, immutable state management. Built on Pydantic v2. Every interaction serializable, auditable, and reproducible.\nTech: Pydantic v2 · Event sourcing · Supabase\n\nLayer 6: Infrastructure (Foundation)\nSupabase (Postgres + pgvector + Auth), Python 3.11+, asyncio, OpenAI/Anthropic/Gemini SDKs, Redis (optional caching)"),
            ("HyperState — Typed State Management", "Every agent interaction is represented as a fully typed, Pydantic v2 validated state object. State transitions are atomic, logged, and replayable. No silent mutations — the system fails loudly if state integrity is violated.\n\nState schema:\nclass AgentState(BaseModel):\n    task_id: UUID\n    agent_id: str\n    status: StateStatus\n    context: dict[str, Any]\n    created_at: datetime\n    checksum: str"),
            ("HyperRouter — UCB1 Bandit Algorithm", "UCB1 (Upper Confidence Bound) is a provably optimal exploration/exploitation algorithm from reinforcement learning. HyperRouter tracks success rates per agent per task type, automatically routing to the highest expected performer while still exploring alternatives.\n\nUCB1 formula:\nUCB1 score = x̄ᵢ + √(2 ln(n) / nᵢ)\n\nWhere:\nx̄ᵢ = average reward for agent i\nn  = total tasks dispatched\nnᵢ = tasks sent to agent i"),
            ("HyperMemory — Causal Graph", "Most vector memory systems store 'what happened.' HyperMemory stores why — causal edges between memory nodes let agents reason about dependencies, timelines, and consequences. Powered by pgvector for sub-millisecond semantic retrieval at scale.\n\nMemory node schema:\n{ id, embedding: vector(1536),\n  content, timestamp, agent_id,\n  causal_parents: [id...],\n  causal_children: [id...] }"),
            ("HyperShield — Policy Engine", "Before any agent executes an action, HyperShield evaluates it against a configurable policy ruleset. Rules can restrict data access, cap API spend, enforce content guidelines, or require human-in-the-loop approval. Full audit log for every decision.\n\nPolicy example:\n@shield.policy('no_pii_export')\nasync def validate(action: Action):\n    if action.contains_pii():\n        raise PolicyViolation('PII export blocked')"),
            ("PROMETHEUS — Recursive Intelligence", "PROMETHEUS is the executive intelligence of HyperClaw. It sits above all subsystems as a meta-agent capable of reasoning about agent selection, spawning new agents, coordinating parallel workflows, and synthesizing results from multi-agent execution trees.\n\nUnlike static orchestration patterns (sequences, DAGs), PROMETHEUS generates its coordination strategy at runtime based on task decomposition. It uses HyperRouter's UCB1 data for agent selection, HyperMemory's causal graph to avoid redundant work, and HyperShield to ensure every spawned agent operates within policy bounds.\n\nKey capabilities: Dynamic agent spawning · Task decomposition · Parallel execution trees · Result synthesis · Self-termination when objectives are met."),
            ("Quality & Test Coverage", "Tests Passing: 377\nKnown Failures: 0\nCore License: MIT\nRelease: v0.1.0-alpha\nTest framework: pytest + asyncio\nCoverage: Unit + integration across all four pillars"),
            ("Full Technology Stack", "Runtime: Python 3.11+ (asyncio-native)\nState Validation: Pydantic v2 (full type safety, JSON schema export)\nDatabase: Supabase/Postgres (RLS, real-time subscriptions)\nVector Memory: pgvector (1536-dim vectors, HNSW index)\nLLM Providers: OpenAI / Anthropic / Gemini (multi-model via HyperRouter)\nCaching: Redis optional (reduces LLM cost on repeated tasks)\nTesting: pytest + asyncio (377 passing)\nPackaging: PyPI (pip install hyperclaw)\nSource: GitHub (github.com/mentatalbans/hyperclaw)"),
        ]
    },
    {
        "path": "06-team/HC-Leadership-Team.html",
        "title": "HyperClaw — Leadership Team",
        "sections": [
            ("HyperClaw Leadership Team", "Founder Bios & Credentials | Q1 2026 | Confidential"),
            ("Team Overview", "HyperClaw is led by a founder pairing of deep enterprise operator experience and AI infrastructure engineering expertise — the combination required to build and sell production AI systems at scale."),
            ("Vaughn Davis — Founder & CEO", "Vaughn Davis brings 15+ years of luxury lifestyle hospitality and Industry 4.0 leadership to HyperClaw. As former Managing Director of Dream Hollywood and CEO of both Hyper Nimbus and Hyper Talent, Vaughn has spent his career at the intersection of high-touch human service and cutting-edge technology adoption.\n\nHe pioneered AI-driven guest experience systems and robotic concierge deployments at scale — making him one of the few enterprise operators who has actually shipped production AI to real customers in high-stakes environments. This operational credibility is HyperClaw's sales moat: Vaughn sells from direct experience, not demos.\n\nCredentials:\n• Former Managing Director — Dream Hollywood Hotel, Dream Hotel Group\n• CEO, Hyper Nimbus — AI-powered hospitality management platform (launched Feb 2024)\n• CEO, Hyper Talent — AI-native talent management division\n• Board Member — Hollywood Chamber of Commerce, The Hollywood Partnership\n• DEI Committee, LA Tourism Board\n• Hofstra University — Political Science\n• Career spanning: Gansevoort Hotel Group, Hyatt, Dream Hotel Group, Two Roads Hospitality"),
            ("Andy Zhang — Co-Founder & CTO", "Andy Zhang is a seasoned AI infrastructure engineer and serial technical co-founder. He previously co-founded Fairsplit, a fintech product built around equitable financial splitting algorithms, and Cogniscale, an AI infrastructure company focused on scalable model deployment.\n\nAt HyperClaw, Andy designed and built the four-pillar architecture — HyperState, HyperRouter, HyperMemory, HyperShield — and the PROMETHEUS recursive intelligence layer. The 377-test codebase under MIT license reflects his engineering rigor. He is the reason HyperClaw ships production-ready systems, not research prototypes.\n\nCredentials:\n• Co-Founder, Cogniscale — AI infrastructure for scalable model deployment\n• Co-Founder, Fairsplit — Fintech platform, equitable financial splitting\n• AI Infrastructure Specialist — multi-model orchestration, vector databases, agentic systems\n• HyperClaw Architect — designed HyperState, HyperRouter (UCB1), HyperMemory, HyperShield, PROMETHEUS"),
            ("Why This Team Wins", "1. Operator + Engineer\nThe rarest pairing in enterprise AI: a CEO who has deployed AI in production at scale, and a CTO who builds the systems from scratch. No gap between vision and execution.\n\n2. Enterprise Access\nVaughn's network spans luxury hospitality, LA commerce, and Industry 4.0 — opening doors to early enterprise adopters that most AI startups cannot reach in their first year.\n\n3. Technical Credibility\nAndy's track record at Cogniscale and Fairsplit, combined with 377 passing tests and a production-grade architecture, signals to enterprise buyers that HyperClaw ships real software.\n\n4. Cross-Industry Vision\nHyperClaw's agent domains span Personal, Business, Scientific, Creative, and Recursive — a breadth of vision that reflects a founding team that understands AI's horizontal potential.\n\n5. Execution Velocity\nv0.1.0-alpha shipped with Supabase backend live, 377 tests, and MIT license. This team ships.\n\n6. Complementary DNA\nVaughn handles go-to-market, enterprise relationships, and vision. Andy owns architecture, engineering quality, and technical roadmap. Clean ownership, no overlap, no gaps."),
        ]
    },
    {
        "path": "07-legal/HC-Risk-Factors.html",
        "title": "HyperClaw — Risk Factors",
        "sections": [
            ("HyperClaw Risk Factors", "Pre-Seed Investment Risk Disclosure | Q1 2026 | Confidential"),
            ("Disclosure Notice", "Investing in early-stage companies involves significant risk. The following risk factors are material to an investment decision in HyperClaw and should be considered carefully alongside the company's financial projections and strategic plan.\n\nThis document contains forward-looking statements and risk disclosures for informational purposes. It does not constitute an offer to sell securities. All projections are estimates based on current assumptions and are subject to change. Investors should consult qualified legal and financial counsel before making investment decisions."),
            ("Risk 1 — Open-Source Competition (HIGH)", "RISK: LangChain, AutoGen, and CrewAI are well-funded, widely adopted open-source frameworks with large communities and brand recognition. Their gravitational pull on developers is significant. Microsoft backs AutoGen; Sequoia has invested in LangChain. A sustained push by any to add enterprise features could compress HyperClaw's differentiation window.\n\nMITIGANTS:\n• HyperShield policy engine requires 12–18 months of enterprise feedback to replicate properly\n• UCB1 routing is architecturally baked in, not a plugin — migration cost is high for incumbents\n• HyperClaw targets enterprise buyers, not developer communities — different motion, different moat\n• MIT core means HyperClaw can become the standard layer on top of any framework"),
            ("Risk 2 — Enterprise Sales Cycle (HIGH)", "RISK: Enterprise software procurement cycles typically run 6–18 months for initial contracts. A small team burning through a $2M raise could face cash constraints before deals close, particularly if early enterprise bets require significant customization.\n\nMITIGANTS:\n• Vaughn's existing network in luxury hospitality and LA enterprise cuts procurement time via champion relationships\n• $499/mo Starter tier enables bottom-up land-and-expand — no 12-month procurement needed to start\n• Open-source core allows IT teams to evaluate HyperClaw before procurement gets involved\n• 18-month runway on $2M raise provides cushion through longer deal cycles"),
            ("Risk 3 — Key Person Risk (MEDIUM)", "RISK: With a two-person founding team, HyperClaw is exposed to key person concentration risk. Loss, incapacitation, or departure of either founder would significantly impact company operations.\n\nMITIGANTS:\n• Pre-seed capital earmarked for 3 engineering hires — distributing technical concentration\n• MIT-licensed codebase with 377 tests is highly documented and transferable\n• Vesting schedules and standard founder protections being established at raise\n• Key-person insurance to be established post-raise"),
            ("Risk 4 — Market Timing (MEDIUM)", "RISK: A contraction in enterprise AI spending, or a major AI incident causing regulatory backlash, could slow adoption. Alternatively, the market could move faster than projected, favoring better-funded competitors.\n\nMITIGANTS:\n• Enterprise AI adoption is a 10-year structural shift, not a single-year cycle\n• HyperShield is more valuable in a regulated environment — regulatory risk is a tailwind\n• Open-source core creates organic adoption independent of sales cycles\n• Revenue model includes low-entry Starter tier ($499/mo)"),
            ("Risk 5 — LLM Provider Dependency (MEDIUM)", "RISK: HyperClaw relies on OpenAI, Anthropic, and Google LLM APIs. Pricing changes, API deprecations, or service outages could impact reliability and unit economics.\n\nMITIGANTS:\n• HyperRouter's multi-model design allows instant failover between providers\n• No hard dependency on any single provider — swappable via configuration\n• Open-source LLM support (Llama, Mistral) on roadmap for on-prem sovereignty tier\n• Sovereign tier explicitly designed for air-gapped, self-hosted model deployment"),
            ("Risk 6 — Runway & Capital Efficiency (LOW)", "RISK: $2M provides approximately 18 months of runway. If ARR targets are not met or the Series A market tightens, HyperClaw could face a bridge financing gap.\n\nMITIGANTS:\n• Supabase backend eliminates significant infrastructure spend at early stage\n• Open-source community creates free distribution and reduces marketing spend\n• Revenue model activates at Starter ($499/mo) — low bar to first revenue\n• Hiring plan tied to ARR milestones, not calendar"),
            ("Risk 7 — AI Regulatory Environment (LOW)", "RISK: The EU AI Act, potential US federal AI regulation, and state-level AI laws could impose compliance requirements. Evolving definitions of 'high-risk AI' could require additional certification.\n\nMITIGANTS:\n• HyperShield policy engine is explicitly designed for compliance requirements — regulation is a tailwind\n• Audit logging built into HyperState from day one\n• MIT license and open architecture make third-party compliance audits straightforward\n• Sovereign tier designed for highest-compliance environments"),
        ]
    },
    {
        "path": "08-traction/HC-Traction-Summary.html",
        "title": "HyperClaw — Traction Summary",
        "sections": [
            ("HyperClaw Traction Summary", "Pre-Seed Raise | Q1 2026 | Confidential"),
            ("Summary", "HyperClaw has shipped a production alpha, established its cloud infrastructure, and validated the core architecture through 377 passing tests. We are raising $2M to accelerate from alpha to first enterprise customers."),
            ("Key Metrics — Q1 2026", "Tests Passing: 377 (zero failures)\nRelease: v0.1.0-alpha (live)\nAgents Shipped: 22 (across 5 domains)\nPre-Seed Raise: $2,000,000 at $10M pre-money\nGitHub: github.com/mentatalbans/hyperclaw (MIT, public)\nBackend: Supabase (live)"),
            ("Development Milestones", "Q4 2025 — Architecture Design & Core Build\nHyperState, HyperRouter (UCB1), HyperMemory (pgvector causal graph), and HyperShield policy engine designed and implemented. PROMETHEUS recursive intelligence layer architected.\n\nQ4 2025 — 377 Tests Written & Passing [LIVE]\nComprehensive test suite covering all four architecture pillars. Zero known failures.\n\nQ1 2026 — v0.1.0-alpha Released [LIVE]\nFirst public alpha release on GitHub under MIT license. 22 specialist agents across 5 domains. Supabase backend live. Open for community adoption.\n\nQ1 2026 — Supabase Backend Live [LIVE]\nProduction cloud infrastructure deployed. Postgres with pgvector, RLS, real-time subscriptions, and Auth layer all live.\n\nQ1 2026 — Pre-Seed Round Opened [IN PROGRESS]\n$2M raise at $10M pre-money valuation. Targeting lead investor close Q2 2026.\n\nQ2 2026 — First Enterprise Customers (Target)\nClose first 3–5 enterprise contracts. Target ARR: $50K–$200K from initial cohort.\n\nQ3 2026 — v0.2 Release (Target)\nSSO/SAML integration, custom policy templates, PROMETHEUS v2, enhanced HyperMemory."),
            ("ARR Growth Trajectory", "2026: $500,000 ARR (first enterprise customers)\n2027: $3,000,000 ARR (GTM scaling, post-Series A)\n2028: $12,000,000 ARR (market expansion)\n2030: $50,000,000 ARR (market leadership)"),
            ("Pre-Seed Raise Terms", "Raise Amount: $2,000,000\nPre-Money Valuation: $10,000,000\nPost-Money Valuation: $12,000,000\nInstrument: SAFE / Priced Round\nTarget Close: Q2 2026\nProjected Runway: ~18 months"),
            ("Use of Funds", "Engineering Hires (40% / $800K): 3 senior engineering hires — backend AI, ML engineer, DevOps\nGTM & Sales (25% / $500K): SDR hire, enterprise sales infrastructure, conference presence\nInfrastructure (20% / $400K): Supabase scaling, LLM API costs, security tooling\nOperations (10% / $200K): HR/recruiting, team travel, equipment\nLegal & Admin (5% / $100K): Corporate legal, IP counsel, option plan documentation"),
            ("Open Source Presence", "Repository: github.com/mentatalbans/hyperclaw\nLicense: MIT (open-source core)\nRelease: v0.1.0-alpha\nTest coverage: 377 passing\nVisibility: Public\n\nThe open-source strategy creates a community-driven top-of-funnel. Developers discover HyperClaw, evaluate it in their organizations, and create the enterprise demand that the sales team converts to paid contracts."),
        ]
    },
]


def set_heading_style(paragraph, level=1):
    """Style heading paragraphs."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
    else:
        run.font.size = Pt(12)
        run.font.bold = True


def create_docx(doc_data):
    """Create a .docx file for a given document."""
    html_path = BASE_DIR / doc_data["path"]
    docx_path = html_path.with_suffix(".docx")
    
    doc = Document()
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Document title
    title_para = doc.add_heading(doc_data["title"], 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Confidentiality notice
    conf_para = doc.add_paragraph("CONFIDENTIAL — HyperClaw AI Orchestration Platform — Q1 2026")
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf_para.runs[0]
    conf_run.font.size = Pt(9)
    conf_run.font.italic = True
    
    doc.add_paragraph()  # spacing
    
    for section_title, section_content in doc_data["sections"]:
        # Section heading
        heading = doc.add_heading(section_title, level=1)
        
        # Section content
        if section_content.strip():
            # Split by double newlines for paragraphs
            paragraphs = section_content.strip().split("\n\n")
            for para_text in paragraphs:
                if para_text.strip():
                    # Check if it's a bullet list
                    lines = para_text.strip().split("\n")
                    if any(line.startswith("•") or line.startswith("·") for line in lines):
                        for line in lines:
                            line = line.strip()
                            if line.startswith("•") or line.startswith("·"):
                                p = doc.add_paragraph(line[1:].strip(), style="List Bullet")
                            elif line:
                                p = doc.add_paragraph(line)
                    else:
                        p = doc.add_paragraph(para_text.strip())
                        p.style = doc.styles["Normal"]
        
        doc.add_paragraph()  # spacing between sections
    
    # Footer note
    doc.add_paragraph()
    footer_para = doc.add_paragraph(
        f"HyperClaw AI Orchestration Platform · {doc_data['title']} · Confidential · Q1 2026\n"
        "GitHub: github.com/mentatalbans/hyperclaw · MIT Licensed Core"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(9)
        run.font.italic = True
    
    doc.save(str(docx_path))
    return docx_path.name


# Generate all docs
print(f"Generating {len(DOCS)} .docx files...")
for doc_data in DOCS:
    try:
        name = create_docx(doc_data)
        print(f"  ✅ {name}")
    except Exception as e:
        print(f"  ❌ {doc_data['path']}: {e}")

print(f"\nDone: {len(DOCS)} DOCX files generated")
